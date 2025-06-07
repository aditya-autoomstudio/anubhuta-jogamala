"""
Word Document OCR Error Analyzer and Cleaner

This module analyzes Word documents for OCR errors and provides intelligent
corrections specifically for Odia medical texts from Anubhuta Jogamala.

Features:
- Extract text from Word documents
- Analyze OCR error patterns
- Generate correction mappings
- Clean and correct misrecognized text
- Preserve document structure while fixing content
"""

import logging
from pathlib import Path
from typing import Dict, List, Tuple, Set
import re
from collections import Counter, defaultdict

# Word document processing
from docx import Document
import json

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WordDocumentCleaner:
    """
    Comprehensive OCR error analyzer and cleaner for Word documents.
    
    Specifically designed for Odia medical texts with common OCR
    misrecognition patterns and intelligent correction suggestions.
    """
    
    def __init__(self) -> None:
        """Initialize the cleaner with OCR error patterns and corrections."""
        self.setup_odia_patterns()
        self.setup_medical_vocabulary()
        self.setup_common_corrections()
        self.analysis_results = {}
        
    def setup_odia_patterns(self) -> None:
        """Setup common Odia OCR error patterns and corrections."""
        
        # Common OCR symbol errors in Odia texts
        self.symbol_corrections = {
            # Currency and mathematical symbols misread as Odia
            '€': '',  # Euro symbol
            '©': '',  # Copyright
            '™': '',  # Trademark
            '®': '',  # Registered
            '†': '',  # Dagger
            '‡': '',  # Double dagger
            '§': '',  # Section
            '¶': '',  # Paragraph
            '•': '',  # Bullet
            '‰': '',  # Per mille
            '‱': '',  # Per ten thousand
            '′': '',  # Prime
            '″': '',  # Double prime
            '‴': '',  # Triple prime
            '‵': '',  # Reversed prime
            '‶': '',  # Reversed double prime
            '‷': '',  # Reversed triple prime
            '‸': '',  # Caret
            '‹': '',  # Single left angle quotation
            '›': '',  # Single right angle quotation
            '‼': '',  # Double exclamation
            '‽': '',  # Interrobang
            '⁇': '',  # Double question mark
            '⁈': '',  # Question exclamation mark
            '⁉': '',  # Exclamation question mark
            '⁎': '',  # Low asterisk
            '⁏': '',  # Reversed semicolon
            '⁐': '',  # Close up
            '⁑': '',  # Two asterisks aligned vertically
            '⁒': '',  # Commercial minus sign
            '⁓': '',  # Swung dash
            '⁔': '',  # Inverted undertie
            '⁕': '',  # Flower punctuation mark
            '⁖': '',  # Three dot punctuation
            '⁗': '',  # Quadruple prime
            '⁘': '',  # Four dot punctuation
            '⁙': '',  # Five dot punctuation
            '⁚': '',  # Two dot punctuation
            '⁛': '',  # Three dot punctuation
            '⁜': '',  # Four dot mark
            '⁝': '',  # Tricolon
            '⁞': '',  # Vertical four dots
            
            # Mathematical symbols
            '×': '',  # Multiplication
            '÷': '',  # Division
            '±': '',  # Plus-minus
            '∞': '',  # Infinity
            '∝': '',  # Proportional to
            '∠': '',  # Angle
            '∡': '',  # Measured angle
            '∢': '',  # Spherical angle
            '∣': '',  # Divides
            '∤': '',  # Does not divide
            '∥': '',  # Parallel to
            '∦': '',  # Not parallel to
            '∧': '',  # Logical and
            '∨': '',  # Logical or
            '∩': '',  # Intersection
            '∪': '',  # Union
            '∫': '',  # Integral
            '∬': '',  # Double integral
            '∭': '',  # Triple integral
            '∮': '',  # Contour integral
            '∯': '',  # Surface integral
            '∰': '',  # Volume integral
            '∱': '',  # Clockwise integral
            '∲': '',  # Clockwise contour integral
            '∳': '',  # Anticlockwise contour integral
            
            # Currency symbols
            '¢': '',  # Cent
            '£': '',  # Pound
            '¤': '',  # Generic currency
            '¥': '',  # Yen
            '₹': '',  # Indian rupee
            '₨': '',  # Rupee
            '₩': '',  # Won
            '₪': '',  # New shekel
            '₫': '',  # Dong
            '€': '',  # Euro
            '₯': '',  # Drachma
            '₰': '',  # Pfennig
            '₱': '',  # Peso
            '₲': '',  # Guarani
            '₳': '',  # Austral
            '₴': '',  # Hryvnia
            '₵': '',  # Cedi
            '₶': '',  # Livre tournois
            '₷': '',  # Spesmilo
            '₸': '',  # Tenge
            '₹': '',  # Indian rupee
            '₺': '',  # Turkish lira
            
            # Common OCR number/letter confusions
            '0': 'ଓ',  # Zero to Odia O
            '1': '୧',  # One to Odia one
            '2': '୨',  # Two to Odia two
            '3': '୩',  # Three to Odia three
            '4': '୪',  # Four to Odia four
            '5': '୫',  # Five to Odia five
            '6': '୬',  # Six to Odia six
            '7': '୭',  # Seven to Odia seven
            '8': '୮',  # Eight to Odia eight
            '9': '୯',  # Nine to Odia nine
        }
        
        # Common character-level corrections
        self.char_corrections = {
            # English letters misread as Odia
            'a': 'ଅ',
            'e': 'ଏ',
            'i': 'ଇ',
            'o': 'ଓ',
            'u': 'ଉ',
            
            # Common OCR errors in Odia characters
            'ର୍': 'ର',  # Halanta corrections
            'ନ୍': 'ନ',
            'ତ୍': 'ତ',
            'ସ୍': 'ସ',
            'କ୍': 'କ',
            'ଲ୍': 'ଲ',
            'ମ୍': 'ମ',
            'ପ୍': 'ପ',
            'ବ୍': 'ବ',
            'ଗ୍': 'ଗ',
            'ଦ୍': 'ଦ',
            'ଯ୍': 'ଯ',
            'ଭ୍': 'ଭ',
            'ଜ୍': 'ଜ',
            'ଧ୍': 'ଧ',
            'ଣ୍': 'ଣ',
            'ଚ୍': 'ଚ',
            'ଛ୍': 'ଛ',
            'ଝ୍': 'ଝ',
            'ଟ୍': 'ଟ',
            'ଠ୍': 'ଠ',
            'ଡ୍': 'ଡ',
            'ଢ୍': 'ଢ',
            'ଫ୍': 'ଫ',
            'ଶ୍': 'ଶ',
            'ଷ୍': 'ଷ',
            'ହ୍': 'ହ',
            'କ୍ଷ୍': 'କ୍ଷ',
            'ଜ୍ଞ୍': 'ଜ୍ଞ',
        }
        
    def setup_medical_vocabulary(self) -> None:
        """Setup common medical terms and their correct spellings."""
        
        self.medical_terms = {
            # Core medical terms
            'ରୋଗ': ['ର୍ରୋଗ', 'ରୋଗ୍', 'ର୍ ଗ', 'ରୋ ଗ'],
            'ଚିକିତ୍ସା': ['ଚିକିତସା', 'ଚିକିତ୍‌ସା', 'ଚିକିତସା'],
            'ଯୋଗ': ['ଯ୍ଗ', 'ଯୋଗ୍', 'ଯୋ ଗ', 'ଯୋଗଗ'],
            'ଔଷଧ': ['ଔଷଦ', 'ଔଷଧ୍', 'ଔଷଢ', 'ଔଷ ଧ'],
            'ଦ୍ରବ୍ୟ': ['ଦ୍ରବୟ', 'ଦ୍ରବ୍ଯ', 'ଦ୍ରବୟ', 'ଦରବୟ'],
            'ପତ୍ର': ['ପତର', 'ପତ୍ର୍', 'ପତ୍‌ର', 'ପତ ର'],
            'ମୂଳ': ['ମୂଲ', 'ମୂଳ୍', 'ମୂ ଳ', 'ମୂଳଳ'],
            'ଫଳ': ['ଫଲ', 'ଫଳ୍', 'ଫ ଳ', 'ଫଳଳ'],
            'ରସ': ['ରସ୍', 'ର ସ', 'ରସସ', 'ରସ'],
            'ଜଳ': ['ଜଲ', 'ଜଳ୍', 'ଜ ଳ', 'ଜଳଳ'],
            'ତେଲ': ['ତେଲ୍', 'ତେ ଲ', 'ତେଲଲ', 'ତେଲ'],
            'ଲଗାଇବ': ['ଲଗାଇବ୍', 'ଲଗାଇବବ', 'ଲଗାଇ ବ', 'ଲଗାଇବ'],
            'ପିଆଇବ': ['ପିଆଇବ୍', 'ପିଆଇବବ', 'ପିଆଇ ବ', 'ପିଆଇବ'],
            'ଖୁଆଇବ': ['ଖୁଆଇବ୍', 'ଖୁଆଇବବ', 'ଖୁଆଇ ବ', 'ଖୁଆଇବ'],
            'ସେର': ['ସେର୍', 'ସେ ର', 'ସେରର', 'ସେର'],
            'ପାନ': ['ପାନ୍', 'ପା ନ', 'ପାନନ', 'ପାନ'],
            'ଟୋପା': ['ଟୋପା୍', 'ଟୋ ପା', 'ଟୋପାା', 'ଟୋପା'],
            
            # Disease categories
            'ଜ୍ବର': ['ଜବର', 'ଜ୍ବର୍', 'ଜ୍‌ବର', 'ଜ ବର'],
            'ଉଦର': ['ଉଦର୍', 'ଉ ଦର', 'ଉଦରର', 'ଉଦର'],
            'ପେଟ': ['ପେଟ୍', 'ପେ ଟ', 'ପେଟଟ', 'ପେଟ'],
            'ମୁଣ୍ଡ': ['ମୁଣଡ', 'ମୁଣ୍ଡ୍', 'ମୁଣ୍‌ଡ', 'ମୁ ଣଡ'],
            'ଆଖି': ['ଆଖି୍', 'ଆ ଖି', 'ଆଖିଇ', 'ଆଖି'],
            'କାନ': ['କାନ୍', 'କା ନ', 'କାନନ', 'କାନ'],
            'ନାକ': ['ନାକ୍', 'ନା କ', 'ନାକକ', 'ନାକ'],
            'ଦାନ୍ତ': ['ଦାନତ', 'ଦାନ୍ତ୍', 'ଦାନ୍‌ତ', 'ଦା ନତ'],
            'ଗଳା': ['ଗଳା୍', 'ଗ ଳା', 'ଗଳାା', 'ଗଳା'],
            'ହାତ': ['ହାତ୍', 'ହା ତ', 'ହାତତ', 'ହାତ'],
            'ପାଦ': ['ପାଦ୍', 'ପା ଦ', 'ପାଦଦ', 'ପାଦ'],
            'ଛାତି': ['ଛାତି୍', 'ଛା ତି', 'ଛାତିଇ', 'ଛାତି'],
            'ପିଠି': ['ପିଠି୍', 'ପି ଠି', 'ପିଠିଇ', 'ପିଠି'],
        }
        
    def setup_common_corrections(self) -> None:
        """Setup patterns for common OCR corrections."""
        
        # Regex patterns for systematic corrections
        self.correction_patterns = [
            # Remove isolated symbols
            (r'[€©™®†‡§¶•‰‱′″‴‵‶‷‸‹›‼‽⁇⁈⁉⁎⁏⁐⁑⁒⁓⁔⁕⁖⁗⁘⁙⁚⁛⁜⁝⁞×÷±∞∝∠∡∢∣∤∥∦∧∨∩∪∫∬∭∮∯∰∱∲∳¢£¤¥₹₨₩₪₫€₯₰₱₲₳₴₵₶₷₸₺]', ''),
            
            # Fix spacing around punctuation
            (r'\s+([।,;:!?])', r'\1'),
            (r'([।,;:!?])\s*([।,;:!?])', r'\1\2'),
            
            # Remove multiple spaces
            (r'\s{2,}', ' '),
            
            # Fix common number patterns
            (r'\b[0-9]+\b', lambda m: self.convert_to_odia_numbers(m.group())),
            
            # Remove standalone English letters
            (r'\b[a-zA-Z]\b', ''),
            
            # Fix broken Odia words (characters separated by spaces)
            (r'([ଅ-ୱ])\s+([ଅ-ୱ])', r'\1\2'),
            
            # Remove halanta at end of words inappropriately
            (r'([ଅ-ୱ])୍\s', r'\1 '),
            
            # Fix double punctuation
            (r'।{2,}', '।'),
            (r',{2,}', ','),
            
            # Remove trailing punctuation duplicates
            (r'([।,;:!?])\1+', r'\1'),
        ]
        
    def convert_to_odia_numbers(self, number_str: str) -> str:
        """Convert English numbers to Odia numbers."""
        number_map = {
            '0': '୦', '1': '୧', '2': '୨', '3': '୩', '4': '୪',
            '5': '୫', '6': '୬', '7': '୭', '8': '୮', '9': '୯'
        }
        return ''.join(number_map.get(char, char) for char in number_str)
    
    def extract_text_from_docx(self, docx_path: str) -> List[str]:
        """
        Extract all text from Word document.
        
        Args:
            docx_path: Path to Word document
            
        Returns:
            List of paragraph texts
        """
        try:
            doc = Document(docx_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
                    
            logger.info(f"Extracted {len(paragraphs)} paragraphs from {docx_path}")
            return paragraphs
            
        except Exception as e:
            logger.error(f"Failed to extract text from {docx_path}: {e}")
            return []
    
    def analyze_ocr_errors(self, text_paragraphs: List[str]) -> Dict:
        """
        Analyze OCR errors in the extracted text.
        
        Args:
            text_paragraphs: List of paragraph texts
            
        Returns:
            Dictionary with error analysis
        """
        analysis = {
            'total_paragraphs': len(text_paragraphs),
            'total_characters': 0,
            'error_symbols': Counter(),
            'suspicious_patterns': Counter(),
            'non_odia_chars': Counter(),
            'correction_suggestions': defaultdict(list),
            'word_frequency': Counter(),
            'potential_errors': []
        }
        
        # Combine all text for analysis
        full_text = ' '.join(text_paragraphs)
        analysis['total_characters'] = len(full_text)
        
        # Count error symbols
        for char in full_text:
            if char in self.symbol_corrections:
                analysis['error_symbols'][char] += 1
        
        # Find non-Odia characters (excluding common punctuation)
        odia_range = range(0x0B00, 0x0B80)  # Odia Unicode range
        common_punct = set('.,;:!?()[]{}"\' \n\t-')
        
        for char in full_text:
            if (ord(char) not in odia_range and 
                char not in common_punct and 
                not char.isdigit()):
                analysis['non_odia_chars'][char] += 1
        
        # Analyze words
        words = re.findall(r'[ଅ-ୱ]+', full_text)
        analysis['word_frequency'] = Counter(words)
        
        # Find potential medical term errors
        for correct_term, variations in self.medical_terms.items():
            for variation in variations:
                if variation in full_text:
                    analysis['correction_suggestions'][variation].append(correct_term)
        
        # Find suspicious patterns
        suspicious_patterns = [
            r'[a-zA-Z]{2,}',  # English words
            r'[0-9]{3,}',     # Long numbers
            r'[€©™®†‡§¶•]+',  # Symbol clusters
            r'[ଅ-ୱ]\s[ଅ-ୱ]',  # Separated Odia characters
            r'୍\s',           # Halanta followed by space
        ]
        
        for pattern in suspicious_patterns:
            matches = re.findall(pattern, full_text)
            for match in matches:
                analysis['suspicious_patterns'][match] += 1
        
        return analysis
    
    def clean_text(self, text: str) -> str:
        """
        Clean OCR errors from text.
        
        Args:
            text: Input text with OCR errors
            
        Returns:
            Cleaned text
        """
        cleaned_text = text
        
        # Apply symbol corrections
        for error_char, correction in self.symbol_corrections.items():
            cleaned_text = cleaned_text.replace(error_char, correction)
        
        # Apply character corrections
        for error_char, correction in self.char_corrections.items():
            cleaned_text = cleaned_text.replace(error_char, correction)
        
        # Apply medical term corrections
        for incorrect_term, correct_terms in self.medical_terms.items():
            if correct_terms:  # Take the first (primary) correction
                for variation in correct_terms:
                    cleaned_text = cleaned_text.replace(variation, incorrect_term)
        
        # Apply regex pattern corrections
        for pattern, replacement in self.correction_patterns:
            if callable(replacement):
                cleaned_text = re.sub(pattern, replacement, cleaned_text)
            else:
                cleaned_text = re.sub(pattern, replacement, cleaned_text)
        
        # Final cleanup
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        
        return cleaned_text
    
    def clean_docx_file(self, input_path: str, output_path: str) -> Dict:
        """
        Clean OCR errors in Word document and save corrected version.
        
        Args:
            input_path: Path to input Word document
            output_path: Path to save cleaned document
            
        Returns:
            Dictionary with cleaning statistics
        """
        try:
            doc = Document(input_path)
            stats = {
                'paragraphs_processed': 0,
                'characters_before': 0,
                'characters_after': 0,
                'corrections_made': 0,
                'paragraphs_changed': 0
            }
            
            for para in doc.paragraphs:
                if para.text.strip():
                    original_text = para.text
                    stats['characters_before'] += len(original_text)
                    
                    # Clean the text
                    cleaned_text = self.clean_text(original_text)
                    stats['characters_after'] += len(cleaned_text)
                    
                    # Update paragraph if changed
                    if cleaned_text != original_text:
                        para.clear()
                        para.add_run(cleaned_text)
                        stats['paragraphs_changed'] += 1
                        stats['corrections_made'] += 1
                    
                    stats['paragraphs_processed'] += 1
            
            # Save cleaned document
            doc.save(output_path)
            logger.info(f"Cleaned document saved to: {output_path}")
            
            return stats
            
        except Exception as e:
            logger.error(f"Failed to clean document {input_path}: {e}")
            return {}
    
    def generate_error_report(self, docx_path: str, output_path: str = None) -> Dict:
        """
        Generate comprehensive error analysis report.
        
        Args:
            docx_path: Path to Word document
            output_path: Path to save report (optional)
            
        Returns:
            Error analysis dictionary
        """
        logger.info(f"Analyzing OCR errors in: {docx_path}")
        
        # Extract text
        paragraphs = self.extract_text_from_docx(docx_path)
        
        if not paragraphs:
            logger.error("No text extracted from document")
            return {}
        
        # Analyze errors
        analysis = self.analyze_ocr_errors(paragraphs)
        
        # Generate report
        report = {
            'document': docx_path,
            'analysis_date': str(Path().absolute()),
            'statistics': {
                'total_paragraphs': analysis['total_paragraphs'],
                'total_characters': analysis['total_characters'],
                'error_symbols_count': len(analysis['error_symbols']),
                'non_odia_chars_count': len(analysis['non_odia_chars']),
                'unique_words': len(analysis['word_frequency']),
                'suspicious_patterns_count': len(analysis['suspicious_patterns'])
            },
            'top_errors': {
                'error_symbols': dict(analysis['error_symbols'].most_common(20)),
                'non_odia_chars': dict(analysis['non_odia_chars'].most_common(20)),
                'suspicious_patterns': dict(analysis['suspicious_patterns'].most_common(20))
            },
            'correction_suggestions': dict(analysis['correction_suggestions']),
            'common_words': dict(analysis['word_frequency'].most_common(50))
        }
        
        # Save report if path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, ensure_ascii=False, indent=2)
            logger.info(f"Error report saved to: {output_path}")
        
        return report


def main():
    """Main function to analyze and clean Word documents."""
    
    print("📄 WORD DOCUMENT OCR ERROR ANALYZER & CLEANER")
    print("="*60)
    print("🔍 Analyzing OCR errors and generating intelligent corrections...")
    print()
    
    # Initialize cleaner
    cleaner = WordDocumentCleaner()
    
    # Define file paths
    word_dir = Path("Word_Documents")
    if not word_dir.exists():
        print(f"❌ Word documents directory not found: {word_dir}")
        return
    
    # Find Word documents
    docx_files = list(word_dir.glob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~$")]  # Exclude temp files
    
    if not docx_files:
        print(f"❌ No Word documents found in: {word_dir}")
        return
    
    print(f"📁 Found {len(docx_files)} Word document(s) to analyze:")
    for docx_file in docx_files:
        file_size = docx_file.stat().st_size / 1024  # KB
        print(f"   📄 {docx_file.name} ({file_size:.1f} KB)")
    print()
    
    # Create output directories
    reports_dir = Path("OCR_Analysis_Reports")
    cleaned_dir = Path("Cleaned_Word_Documents")
    reports_dir.mkdir(exist_ok=True)
    cleaned_dir.mkdir(exist_ok=True)
    
    # Process each document
    for i, docx_file in enumerate(docx_files, 1):
        print(f"🔍 Analyzing file {i}/{len(docx_files)}: {docx_file.name}")
        
        # Generate error report
        report_path = reports_dir / f"{docx_file.stem}_error_analysis.json"
        report = cleaner.generate_error_report(str(docx_file), str(report_path))
        
        if report:
            print(f"   📊 Analysis complete:")
            print(f"      • Paragraphs: {report['statistics']['total_paragraphs']}")
            print(f"      • Characters: {report['statistics']['total_characters']:,}")
            print(f"      • Error symbols: {report['statistics']['error_symbols_count']}")
            print(f"      • Non-Odia chars: {report['statistics']['non_odia_chars_count']}")
            print(f"      • Suspicious patterns: {report['statistics']['suspicious_patterns_count']}")
            
            # Show top errors
            if report['top_errors']['error_symbols']:
                print(f"      • Top error symbols: {list(report['top_errors']['error_symbols'].keys())[:5]}")
            
        # Clean the document
        cleaned_path = cleaned_dir / f"{docx_file.stem}_cleaned.docx"
        print(f"   🧹 Cleaning document...")
        
        clean_stats = cleaner.clean_docx_file(str(docx_file), str(cleaned_path))
        
        if clean_stats:
            print(f"   ✅ Cleaning complete:")
            print(f"      • Paragraphs processed: {clean_stats['paragraphs_processed']}")
            print(f"      • Paragraphs changed: {clean_stats['paragraphs_changed']}")
            print(f"      • Corrections made: {clean_stats['corrections_made']}")
            
            # Calculate reduction in character count (removed symbols)
            char_reduction = clean_stats['characters_before'] - clean_stats['characters_after']
            if char_reduction > 0:
                print(f"      • Characters removed: {char_reduction:,}")
        
        print()
    
    print("="*60)
    print("🎉 ANALYSIS AND CLEANING COMPLETED!")
    print()
    print(f"📋 GENERATED FILES:")
    print(f"   📊 Error reports: {reports_dir}")
    print(f"   📄 Cleaned documents: {cleaned_dir}")
    print()
    print(f"💡 NEXT STEPS:")
    print(f"   1. Review error analysis reports (JSON files)")
    print(f"   2. Open cleaned Word documents to verify quality")
    print(f"   3. Use correction suggestions for manual refinement")
    print(f"   4. Apply additional domain-specific corrections")
    print(f"   5. Compare original vs cleaned documents")
    print()
    print(f"🔧 RECOMMENDATIONS:")
    print(f"   • Focus on high-frequency error patterns first")
    print(f"   • Use Word's Find & Replace for systematic corrections")
    print(f"   • Verify medical terminology corrections manually")
    print(f"   • Set language to Odia in Word for better spell checking")


if __name__ == "__main__":
    main() 