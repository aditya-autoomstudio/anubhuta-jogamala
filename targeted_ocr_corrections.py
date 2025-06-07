"""
Targeted OCR Corrections for Anubhuta Jogamala Word Documents

This script applies specific corrections based on the error pattern analysis.
It focuses on the most critical and systematic OCR errors identified.
"""

import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Tuple
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class TargetedOCRCorrector:
    """Targeted corrector for specific OCR error patterns."""
    
    def __init__(self):
        """Initialize with specific correction rules."""
        self.setup_correction_rules()
        
    def setup_correction_rules(self) -> None:
        """Setup specific correction rules based on error analysis."""
        
        # Priority 1: Critical character replacements (high frequency errors)
        self.character_replacements = {
            # English to Odia character corrections
            'a': 'ଅ',  # 2326 occurrences in Part 1
            'e': 'ଏ',  # 2623 occurrences in Part 1
            'i': 'ଇ',  # 1314 occurrences in Part 1
            'o': 'ଓ',  # 1170 occurrences
            't': 'ତ',  # 1629 occurrences
            'A': 'ଅ',  # 1484 occurrences
            'n': 'ନ',  # 1066 occurrences
            'r': 'ର',  # 1054 occurrences
            's': 'ସ',  # 1010 occurrences
            'd': 'ଦ',  # 994 occurrences
            'l': 'ଲ',  # 805 occurrences
            'c': 'ଚ',  # 770 occurrences
            'G': 'ଗ',  # 746 occurrences
            'g': 'ଗ',  # 714 occurrences
            'E': 'ଏ',  # 656 occurrences
            
            # Number corrections
            '0': '୦',
            '1': '୧',
            '2': '୨',
            '3': '୩',
            '4': '୪',
            '5': '୫',
            '6': '୬',
            '7': '୭',
            '8': '୮',
            '9': '୯',
            
            # Symbol corrections
            '|': '।',  # Replace pipe with danda
        }
        
        # Priority 2: Pattern-based corrections (regex patterns)
        self.pattern_corrections = [
            # Fix broken words (highest priority - 48,932 errors in Part 1)
            (r'([ଅ-ୱ])\s+([ଅ-ୱାିୀୁୂୃେୈୋୌ])', r'\1\2', 'Join separated Odia characters'),
            
            # Fix incorrect halanta usage (14,031 errors)
            (r'([ଅ-ୱ])୍\s+', r'\1 ', 'Remove trailing halanta'),
            
            # Fix spacing around punctuation (1,682 errors)
            (r'\s+([।,;:!?])', r'\1', 'Remove space before punctuation'),
            (r'([।,;:!?])\s*([।,;:!?])', r'\1\2', 'Fix multiple punctuation'),
            
            # Fix multiple spacing (2,992 errors)
            (r'\s{2,}', ' ', 'Replace multiple spaces with single space'),
            
            # Remove common OCR symbols
            (r'[€©™®†‡§¶•‰‱′″‴‵‶‷‸‹›‼‽⁇⁈⁉⁎⁏⁐⁑⁒⁓⁔⁕⁖⁗⁘⁙⁚⁛⁜⁝⁞×÷±∞∝∠∡∢∣∤∥∦∧∨∩∪∫∬∭∮∯∰∱∲∳¢£¤¥₹₨₩₪₫€₯₰₱₲₳₴₵₶₷₸₺]', '', 'Remove OCR symbol artifacts'),
            
            # Fix underscore artifacts (17,399 occurrences)
            (r'_{2,}', '', 'Remove multiple underscores'),
            (r'_([ଅ-ୱ])', r'\1', 'Remove underscore before Odia character'),
            (r'([ଅ-ୱ])_', r'\1', 'Remove underscore after Odia character'),
            
            # Fix double punctuation
            (r'।{2,}', '।', 'Fix multiple danda'),
            (r',{2,}', ',', 'Fix multiple commas'),
        ]
        
        # Priority 3: Medical term corrections
        self.medical_term_corrections = {
            # Common medical terms with corrections
            'ରୋଗ୍': 'ରୋଗ',
            'ଯୋଗ୍': 'ଯୋଗ',
            'ଚିକିତସା': 'ଚିକିତ୍ସା',
            'ଔଷଧ୍': 'ଔଷଧ',
            'ପତର': 'ପତ୍ର',
            'ମୂଲ': 'ମୂଳ',
            'ଫଳ୍': 'ଫଳ',
            'ରସ୍': 'ରସ',
            'ତେଲ୍': 'ତେଲ',
            'ଜ୍ବର୍': 'ଜ୍ବର',
            'ପେଟ୍': 'ପେଟ',
            'ମୁଣଡ': 'ମୁଣ୍ଡ',
            'ଆଖି୍': 'ଆଖି',
            'କାନ୍': 'କାନ',
            'ନାକ୍': 'ନାକ',
            'ଗଳା୍': 'ଗଳା',
            'ହାତ୍': 'ହାତ',
            'ପାଦ୍': 'ପାଦ',
            'ସେର୍': 'ସେର',
            'ପାନ୍': 'ପାନ',
            'ଟୋପା୍': 'ଟୋପା',
        }
        
    def apply_character_corrections(self, text: str) -> Tuple[str, int]:
        """Apply character-level corrections."""
        corrected_text = text
        corrections_count = 0
        
        for incorrect_char, correct_char in self.character_replacements.items():
            # Only replace standalone characters to avoid corrupting valid text
            if incorrect_char in ['|']:  # Special handling for pipe
                old_count = corrected_text.count(incorrect_char)
                corrected_text = corrected_text.replace(incorrect_char, correct_char)
                corrections_count += old_count
            elif incorrect_char.isalnum():  # Letters and numbers
                # Replace only when surrounded by non-Odia characters or at word boundaries
                pattern = r'\b' + re.escape(incorrect_char) + r'\b'
                matches = len(re.findall(pattern, corrected_text))
                corrected_text = re.sub(pattern, correct_char, corrected_text)
                corrections_count += matches
        
        return corrected_text, corrections_count
    
    def apply_pattern_corrections(self, text: str) -> Tuple[str, int]:
        """Apply pattern-based corrections."""
        corrected_text = text
        corrections_count = 0
        
        for pattern, replacement, description in self.pattern_corrections:
            matches = len(re.findall(pattern, corrected_text))
            corrected_text = re.sub(pattern, replacement, corrected_text)
            corrections_count += matches
            
            if matches > 0:
                logger.info(f"Applied {description}: {matches} corrections")
        
        return corrected_text, corrections_count
    
    def apply_medical_term_corrections(self, text: str) -> Tuple[str, int]:
        """Apply medical terminology corrections."""
        corrected_text = text
        corrections_count = 0
        
        for incorrect_term, correct_term in self.medical_term_corrections.items():
            old_count = corrected_text.count(incorrect_term)
            corrected_text = corrected_text.replace(incorrect_term, correct_term)
            corrections_count += old_count
        
        return corrected_text, corrections_count
    
    def correct_document(self, input_path: str, output_path: str) -> Dict:
        """Apply all corrections to a Word document."""
        
        logger.info(f"Starting corrections on: {input_path}")
        
        try:
            doc = Document(input_path)
        except Exception as e:
            logger.error(f"Failed to open document: {e}")
            return {}
        
        stats = {
            'paragraphs_processed': 0,
            'paragraphs_changed': 0,
            'total_corrections': 0,
            'character_corrections': 0,
            'pattern_corrections': 0,
            'medical_corrections': 0,
            'original_characters': 0,
            'final_characters': 0
        }
        
        for para in doc.paragraphs:
            if para.text.strip():
                original_text = para.text
                stats['original_characters'] += len(original_text)
                
                # Apply corrections in order of priority
                corrected_text = original_text
                
                # 1. Character corrections
                corrected_text, char_count = self.apply_character_corrections(corrected_text)
                stats['character_corrections'] += char_count
                
                # 2. Pattern corrections
                corrected_text, pattern_count = self.apply_pattern_corrections(corrected_text)
                stats['pattern_corrections'] += pattern_count
                
                # 3. Medical term corrections
                corrected_text, medical_count = self.apply_medical_term_corrections(corrected_text)
                stats['medical_corrections'] += medical_count
                
                # Final cleanup
                corrected_text = re.sub(r'\s+', ' ', corrected_text).strip()
                stats['final_characters'] += len(corrected_text)
                
                # Update paragraph if changed
                if corrected_text != original_text:
                    para.clear()
                    para.add_run(corrected_text)
                    stats['paragraphs_changed'] += 1
                
                stats['paragraphs_processed'] += 1
        
        # Calculate total corrections
        stats['total_corrections'] = (stats['character_corrections'] + 
                                    stats['pattern_corrections'] + 
                                    stats['medical_corrections'])
        
        # Save corrected document
        try:
            doc.save(output_path)
            logger.info(f"Corrected document saved to: {output_path}")
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return {}
        
        return stats
    
    def process_all_documents(self) -> None:
        """Process all Word documents in the Word_Documents directory."""
        
        word_dir = Path("Word_Documents")
        if not word_dir.exists():
            logger.error(f"Word documents directory not found: {word_dir}")
            return
        
        # Find Word documents
        docx_files = list(word_dir.glob("*.docx"))
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        
        if not docx_files:
            logger.error(f"No Word documents found in: {word_dir}")
            return
        
        # Create output directory
        output_dir = Path("Targeted_Corrected_Documents")
        output_dir.mkdir(exist_ok=True)
        
        print(f"🔧 TARGETED OCR CORRECTIONS")
        print("="*50)
        print(f"📁 Found {len(docx_files)} document(s) to correct:")
        for docx_file in docx_files:
            print(f"   📄 {docx_file.name}")
        print()
        
        # Process each document
        total_stats = {
            'documents_processed': 0,
            'total_paragraphs': 0,
            'total_corrections': 0,
            'character_corrections': 0,
            'pattern_corrections': 0,
            'medical_corrections': 0
        }
        
        for i, docx_file in enumerate(docx_files, 1):
            print(f"🔧 Processing {i}/{len(docx_files)}: {docx_file.name}")
            
            output_path = output_dir / f"{docx_file.stem}_targeted_corrected.docx"
            stats = self.correct_document(str(docx_file), str(output_path))
            
            if stats:
                print(f"   ✅ Corrections applied:")
                print(f"      • Paragraphs processed: {stats['paragraphs_processed']}")
                print(f"      • Paragraphs changed: {stats['paragraphs_changed']}")
                print(f"      • Total corrections: {stats['total_corrections']:,}")
                print(f"      • Character corrections: {stats['character_corrections']:,}")
                print(f"      • Pattern corrections: {stats['pattern_corrections']:,}")
                print(f"      • Medical term corrections: {stats['medical_corrections']:,}")
                
                # Character reduction
                char_reduction = stats['original_characters'] - stats['final_characters']
                if char_reduction > 0:
                    print(f"      • Characters removed: {char_reduction:,}")
                
                # Update totals
                total_stats['documents_processed'] += 1
                total_stats['total_paragraphs'] += stats['paragraphs_processed']
                total_stats['total_corrections'] += stats['total_corrections']
                total_stats['character_corrections'] += stats['character_corrections']
                total_stats['pattern_corrections'] += stats['pattern_corrections']
                total_stats['medical_corrections'] += stats['medical_corrections']
            
            print()
        
        # Display summary
        print("="*50)
        print("🎉 TARGETED CORRECTIONS COMPLETED!")
        print()
        print(f"📊 SUMMARY:")
        print(f"   • Documents processed: {total_stats['documents_processed']}")
        print(f"   • Total paragraphs: {total_stats['total_paragraphs']}")
        print(f"   • Total corrections: {total_stats['total_corrections']:,}")
        print(f"   • Character corrections: {total_stats['character_corrections']:,}")
        print(f"   • Pattern corrections: {total_stats['pattern_corrections']:,}")
        print(f"   • Medical corrections: {total_stats['medical_corrections']:,}")
        print()
        print(f"📂 Corrected documents saved in: {output_dir}")
        print()
        print(f"💡 NEXT STEPS:")
        print(f"   1. Review corrected documents for quality")
        print(f"   2. Compare with original documents")
        print(f"   3. Apply additional manual corrections if needed")
        print(f"   4. Use corrected versions for further processing")


def main():
    """Main function to run targeted OCR corrections."""
    
    corrector = TargetedOCRCorrector()
    corrector.process_all_documents()


if __name__ == "__main__":
    main() 