"""
PDF to Word Converter for Anubhuta Jogamala Project

This module converts PDF files to Word documents while preserving:
- Tabular data and structure
- Original formatting and layout
- High-quality OCR for Odia text
- Images and graphical elements
- Page breaks and sections
"""

import logging
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import time

# PDF processing libraries
import fitz  # PyMuPDF for high-quality PDF processing
from pdf2image import convert_from_path
import pytesseract
from PIL import Image, ImageEnhance

# Word document creation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class PDFToWordConverter:
    """
    High-quality PDF to Word converter with OCR support for Odia text.
    
    Preserves document structure, tables, and formatting while providing
    clean text extraction and Word document generation.
    """
    
    def __init__(self) -> None:
        """Initialize the converter with optimal settings."""
        self.setup_ocr_config()
        self.setup_conversion_settings()
        self.converted_files = 0
        self.total_pages_processed = 0
        
    def setup_ocr_config(self) -> None:
        """Configure OCR settings for optimal Odia text recognition."""
        
        # Tesseract configuration for Odia + English + Hindi
        self.ocr_config = {
            'lang': 'ori+eng+hin',  # Odia + English + Hindi
            'config': '--oem 3 --psm 6 -c preserve_interword_spaces=1'
        }
        
        # Image preprocessing settings for better OCR
        self.image_settings = {
            'dpi': 300,  # High DPI for quality
            'format': 'PNG',  # Best quality format
            'enhance_contrast': True,
            'sharpen': True,
            'denoise': True
        }
        
    def setup_conversion_settings(self) -> None:
        """Configure conversion settings for optimal results."""
        
        self.conversion_settings = {
            'preserve_tables': True,
            'preserve_images': True,
            'preserve_formatting': True,
            'extract_text_blocks': True,
            'merge_text_blocks': True,
            'min_font_size': 6,  # Minimum font size to process
            'max_font_size': 72,  # Maximum font size to process
        }
        
    def enhance_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Enhance image quality for better OCR results.
        
        Args:
            image: PIL Image to enhance
            
        Returns:
            Enhanced PIL Image
        """
        try:
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Enhance contrast
            if self.image_settings['enhance_contrast']:
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(1.2)
            
            # Enhance sharpness
            if self.image_settings['sharpen']:
                enhancer = ImageEnhance.Sharpness(image)
                image = enhancer.enhance(1.1)
            
            # Convert to grayscale for better OCR
            image = image.convert('L')
            
            return image
            
        except Exception as e:
            logger.warning(f"Image enhancement failed: {e}")
            return image
    
    def extract_text_with_ocr(self, image: Image.Image) -> str:
        """
        Extract text from image using OCR.
        
        Args:
            image: PIL Image to process
            
        Returns:
            Extracted text string
        """
        try:
            # Enhance image for better OCR
            enhanced_image = self.enhance_image_for_ocr(image)
            
            # Perform OCR
            text = pytesseract.image_to_string(
                enhanced_image,
                lang=self.ocr_config['lang'],
                config=self.ocr_config['config']
            )
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""
    
    def extract_pdf_content(self, pdf_path: str) -> List[Dict]:
        """
        Extract content from PDF including text, images, and structure.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of page content dictionaries
        """
        pages_content = []
        
        try:
            # Open PDF with PyMuPDF for structure analysis
            pdf_document = fitz.open(pdf_path)
            
            logger.info(f"Processing PDF: {pdf_path} ({len(pdf_document)} pages)")
            
            # Convert PDF to high-quality images
            images = convert_from_path(
                pdf_path,
                dpi=self.image_settings['dpi'],
                fmt=self.image_settings['format']
            )
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_image = images[page_num]
                
                logger.info(f"Processing page {page_num + 1}/{len(pdf_document)}")
                
                page_content = {
                    'page_number': page_num + 1,
                    'text_blocks': [],
                    'images': [],
                    'tables': [],
                    'ocr_text': '',
                    'page_size': page.rect
                }
                
                # Extract text blocks with positioning
                text_blocks = page.get_text("dict")
                if text_blocks and 'blocks' in text_blocks:
                    for block in text_blocks['blocks']:
                        if 'lines' in block:  # Text block
                            block_text = ""
                            for line in block['lines']:
                                line_text = ""
                                for span in line['spans']:
                                    line_text += span['text']
                                block_text += line_text + "\n"
                            
                            if block_text.strip():
                                page_content['text_blocks'].append({
                                    'text': block_text.strip(),
                                    'bbox': block['bbox'],
                                    'type': 'text'
                                })
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        if pix.n - pix.alpha < 4:  # Skip if not RGB/Gray
                            img_data = pix.tobytes("png")
                            page_content['images'].append({
                                'data': img_data,
                                'index': img_index,
                                'xref': xref
                            })
                        pix = None
                    except Exception as e:
                        logger.warning(f"Could not extract image {img_index}: {e}")
                
                # Perform OCR on the entire page
                page_content['ocr_text'] = self.extract_text_with_ocr(page_image)
                
                # Try to detect tables (basic detection)
                try:
                    tables = page.find_tables()
                    if tables:
                        for table in tables:
                            try:
                                table_data = table.extract()
                                page_content['tables'].append({
                                    'data': table_data,
                                    'bbox': table.bbox
                                })
                            except Exception as e:
                                logger.warning(f"Could not extract table: {e}")
                except Exception as e:
                    logger.debug(f"Table detection not available: {e}")
                
                pages_content.append(page_content)
                self.total_pages_processed += 1
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            
        return pages_content
    
    def create_word_document(self, pages_content: List[Dict], output_path: str) -> bool:
        """
        Create Word document from extracted PDF content.
        
        Args:
            pages_content: List of page content dictionaries
            output_path: Path for output Word document
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create new Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = "Anubhuta Jogamala - Digitized Content"
            doc.core_properties.author = "PDF to Word Converter"
            doc.core_properties.subject = "Traditional Odia Medical Text"
            
            for page_num, page_content in enumerate(pages_content):
                logger.info(f"Creating Word content for page {page_num + 1}")
                
                # Add page header
                if page_num > 0:
                    doc.add_page_break()
                
                # Add page number heading
                heading = doc.add_heading(f"Page {page_content['page_number']}", level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add tables first (if any)
                if page_content.get('tables'):
                    table_heading = doc.add_heading("Tabular Data", level=3)
                    for table_index, table_info in enumerate(page_content['tables']):
                        try:
                            table_data = table_info['data']
                            if table_data and len(table_data) > 0:
                                # Create Word table
                                word_table = doc.add_table(
                                    rows=len(table_data), 
                                    cols=max(len(row) for row in table_data) if table_data else 1
                                )
                                word_table.style = 'Table Grid'
                                
                                # Fill table data
                                for row_idx, row_data in enumerate(table_data):
                                    for col_idx, cell_data in enumerate(row_data):
                                        if row_idx < len(word_table.rows) and col_idx < len(word_table.rows[row_idx].cells):
                                            cell = word_table.rows[row_idx].cells[col_idx]
                                            cell.text = str(cell_data) if cell_data else ""
                                
                                doc.add_paragraph()  # Add space after table
                        except Exception as e:
                            logger.warning(f"Could not add table {table_index}: {e}")
                
                # Add structured text blocks
                if page_content.get('text_blocks'):
                    text_heading = doc.add_heading("Structured Text", level=3)
                    for block in page_content['text_blocks']:
                        if block['text'].strip():
                            para = doc.add_paragraph(block['text'])
                            # Set Odia-friendly font
                            for run in para.runs:
                                run.font.name = 'Kalinga'  # Good Odia font
                                run.font.size = Pt(12)
                
                # Add OCR extracted text
                if page_content.get('ocr_text') and page_content['ocr_text'].strip():
                    ocr_heading = doc.add_heading("OCR Extracted Text", level=3)
                    ocr_para = doc.add_paragraph(page_content['ocr_text'])
                    # Set Odia-friendly font for OCR text
                    for run in ocr_para.runs:
                        run.font.name = 'Kalinga'
                        run.font.size = Pt(11)
                
                # Add separator
                doc.add_paragraph("_" * 80)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Word document saved: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            return False
    
    def convert_pdf_to_word(self, pdf_path: str, output_dir: str) -> bool:
        """
        Convert a single PDF to Word document.
        
        Args:
            pdf_path: Path to input PDF file
            output_dir: Directory for output Word file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            pdf_file = Path(pdf_path)
            if not pdf_file.exists():
                logger.error(f"PDF file not found: {pdf_path}")
                return False
            
            # Create output directory
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Generate output filename
            output_filename = pdf_file.stem + ".docx"
            output_file_path = output_path / output_filename
            
            logger.info(f"Converting: {pdf_path} -> {output_file_path}")
            
            # Extract PDF content
            pages_content = self.extract_pdf_content(pdf_path)
            
            if not pages_content:
                logger.error(f"No content extracted from PDF: {pdf_path}")
                return False
            
            # Create Word document
            success = self.create_word_document(pages_content, str(output_file_path))
            
            if success:
                self.converted_files += 1
                logger.info(f"Successfully converted: {pdf_path}")
                return True
            else:
                logger.error(f"Failed to create Word document for: {pdf_path}")
                return False
                
        except Exception as e:
            logger.error(f"Conversion failed for {pdf_path}: {e}")
            return False
    
    def convert_directory(
        self, 
        input_dir: str, 
        output_dir: str, 
        file_pattern: str = "*.pdf"
    ) -> Dict[str, int]:
        """
        Convert all PDF files in a directory to Word documents.
        
        Args:
            input_dir: Input directory containing PDF files
            output_dir: Output directory for Word files
            file_pattern: File pattern to match
            
        Returns:
            Dictionary with conversion statistics
        """
        input_path = Path(input_dir)
        
        if not input_path.exists():
            logger.error(f"Input directory not found: {input_dir}")
            return {"error": "Input directory not found"}
        
        # Find all PDF files
        pdf_files = list(input_path.glob(file_pattern))
        
        if not pdf_files:
            logger.warning(f"No PDF files found in {input_dir}")
            return {"files_found": 0, "files_converted": 0, "pages_processed": 0}
        
        logger.info(f"Found {len(pdf_files)} PDF files to convert")
        
        # Reset counters
        self.converted_files = 0
        self.total_pages_processed = 0
        
        successful_conversions = 0
        failed_conversions = 0
        
        start_time = time.time()
        
        for pdf_file in pdf_files:
            if self.convert_pdf_to_word(str(pdf_file), output_dir):
                successful_conversions += 1
            else:
                failed_conversions += 1
        
        end_time = time.time()
        
        stats = {
            "files_found": len(pdf_files),
            "files_converted": successful_conversions,
            "files_failed": failed_conversions,
            "pages_processed": self.total_pages_processed,
            "processing_time": round(end_time - start_time, 2)
        }
        
        logger.info(f"Conversion completed: {stats}")
        return stats


def main() -> None:
    """Main function to convert PDFs to Word documents."""
    
    print("📄 PDF TO WORD CONVERTER - ANUBHUTA JOGAMALA PROJECT")
    print("="*70)
    print("🔄 Converting PDF files to Word documents with OCR and structure preservation...")
    print()
    
    # Initialize converter
    converter = PDFToWordConverter()
    
    # Define directories
    input_dir = "Part_1"  # Directory containing Part 1 PDFs
    output_dir = "Part_1_Word"  # Output directory for Word files
    
    print(f"📂 Input directory: {input_dir}")
    print(f"📂 Output directory: {output_dir}")
    print(f"🔧 OCR Languages: {converter.ocr_config['lang']}")
    print(f"🖼️  Image DPI: {converter.image_settings['dpi']}")
    print()
    
    # Convert all PDF files
    stats = converter.convert_directory(input_dir, output_dir)
    
    if "error" in stats:
        print(f"❌ Error: {stats['error']}")
        return
    
    # Display results
    print("✅ CONVERSION RESULTS:")
    print(f"   • PDF files found: {stats['files_found']}")
    print(f"   • Successfully converted: {stats['files_converted']}")
    print(f"   • Failed conversions: {stats['files_failed']}")
    print(f"   • Total pages processed: {stats['pages_processed']}")
    print(f"   • Processing time: {stats['processing_time']} seconds")
    
    if stats['files_converted'] > 0:
        avg_time = stats['processing_time'] / stats['files_converted']
        print(f"   • Average time per file: {avg_time:.1f} seconds")
        
        print(f"\n🎉 Successfully converted {stats['files_converted']} PDF files to Word!")
        print(f"📁 Word documents are available in: {output_dir}")
        
        # Show file list
        output_path = Path(output_dir)
        if output_path.exists():
            word_files = list(output_path.glob("*.docx"))
            if word_files:
                print(f"\n📋 Generated Word files:")
                for word_file in sorted(word_files)[:10]:  # Show first 10
                    print(f"   📄 {word_file.name}")
                if len(word_files) > 10:
                    print(f"   ... and {len(word_files) - 10} more files")
    else:
        print("❌ No files were successfully converted!")
    
    print("\n" + "="*70)
    print("🔧 PDF to Word conversion completed!")
    
    # Recommendations
    if stats['files_converted'] > 0:
        print("\n💡 NEXT STEPS:")
        print("   • Review Word documents for formatting accuracy")
        print("   • Verify OCR text quality for Odia content")
        print("   • Check table preservation and structure")
        print("   • Use Word's built-in tools for further editing")
        print("   • Consider manual review of complex pages")


if __name__ == "__main__":
    main() 